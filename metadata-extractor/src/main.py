"""Metadata Extractor.

A Python script that extracts and reports file metadata including EXIF data
for images, ID3 tags for audio, and document properties for office files.
"""

import argparse
import json
import logging
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

import yaml

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("logs/extractor.log"),
        logging.StreamHandler(sys.stdout),
    ],
)

logger = logging.getLogger(__name__)

try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("PIL/Pillow not available. Image EXIF extraction disabled.")

try:
    from mutagen import File as MutagenFile
    from mutagen.id3 import ID3NoHeaderError
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False
    logger.warning("mutagen not available. Audio ID3 extraction disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX extraction disabled.")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl not available. XLSX extraction disabled.")


class MetadataExtractor:
    """Extracts metadata from various file types."""

    IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".gif", ".webp"}
    AUDIO_EXTENSIONS = {".mp3", ".flac", ".ogg", ".m4a", ".aac", ".wma"}
    OFFICE_EXTENSIONS = {".docx", ".xlsx", ".pptx"}

    def __init__(self) -> None:
        """Initialize the metadata extractor."""
        self.stats = {
            "processed": 0,
            "images": 0,
            "audio": 0,
            "office": 0,
            "errors": 0,
        }

    def extract_image_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract EXIF and image metadata.

        Args:
            file_path: Path to image file

        Returns:
            Dictionary containing image metadata

        Raises:
            ImportError: If PIL/Pillow is not available
        """
        if not PIL_AVAILABLE:
            raise ImportError("PIL/Pillow is required for image metadata extraction")

        metadata: Dict[str, Any] = {
            "file_path": str(file_path),
            "file_type": "image",
            "file_size": file_path.stat().st_size,
        }

        try:
            with Image.open(file_path) as img:
                metadata["format"] = img.format
                metadata["mode"] = img.mode
                metadata["size"] = {"width": img.width, "height": img.height}

                exif_data = img._getexif()
                if exif_data:
                    exif_dict = {}
                    for tag_id, value in exif_data.items():
                        tag = TAGS.get(tag_id, tag_id)
                        exif_dict[tag] = value

                        if tag == "GPSInfo":
                            gps_data = {}
                            for gps_tag_id, gps_value in value.items():
                                gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                                gps_data[gps_tag] = gps_value
                            exif_dict["GPSInfo"] = gps_data

                    metadata["exif"] = exif_dict

                if hasattr(img, "info"):
                    metadata["info"] = dict(img.info)

        except Exception as e:
            logger.warning(f"Error extracting image metadata from {file_path}: {e}")
            metadata["error"] = str(e)

        return metadata

    def extract_audio_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract ID3 tags and audio metadata.

        Args:
            file_path: Path to audio file

        Returns:
            Dictionary containing audio metadata

        Raises:
            ImportError: If mutagen is not available
        """
        if not MUTAGEN_AVAILABLE:
            raise ImportError("mutagen is required for audio metadata extraction")

        metadata: Dict[str, Any] = {
            "file_path": str(file_path),
            "file_type": "audio",
            "file_size": file_path.stat().st_size,
        }

        try:
            audio_file = MutagenFile(file_path)

            if audio_file is None:
                metadata["error"] = "Unable to read audio file"
                return metadata

            metadata["mime_type"] = audio_file.mime[0] if audio_file.mime else None

            if hasattr(audio_file, "info"):
                info = audio_file.info
                if hasattr(info, "length"):
                    metadata["duration"] = info.length
                if hasattr(info, "bitrate"):
                    metadata["bitrate"] = info.bitrate
                if hasattr(info, "sample_rate"):
                    metadata["sample_rate"] = info.sample_rate
                if hasattr(info, "channels"):
                    metadata["channels"] = info.channels

            tags = {}
            for key, value in audio_file.items():
                if isinstance(value, list) and len(value) > 0:
                    tags[key] = value[0] if len(value) == 1 else value
                else:
                    tags[key] = value

            metadata["tags"] = tags

        except ID3NoHeaderError:
            metadata["tags"] = {}
            metadata["note"] = "No ID3 tags found"
        except Exception as e:
            logger.warning(f"Error extracting audio metadata from {file_path}: {e}")
            metadata["error"] = str(e)

        return metadata

    def extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract document properties from DOCX files.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary containing document metadata

        Raises:
            ImportError: If python-docx is not available
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX metadata extraction")

        metadata: Dict[str, Any] = {
            "file_path": str(file_path),
            "file_type": "office_document",
            "file_size": file_path.stat().st_size,
        }

        try:
            doc = DocxDocument(file_path)

            core_props = doc.core_properties
            metadata["properties"] = {
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "comments": core_props.comments,
                "category": core_props.category,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "last_modified_by": core_props.last_modified_by,
                "revision": core_props.revision,
                "language": core_props.language,
            }

            metadata["document_info"] = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

        except Exception as e:
            logger.warning(f"Error extracting DOCX metadata from {file_path}: {e}")
            metadata["error"] = str(e)

        return metadata

    def extract_xlsx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract document properties from XLSX files.

        Args:
            file_path: Path to XLSX file

        Returns:
            Dictionary containing spreadsheet metadata

        Raises:
            ImportError: If openpyxl is not available
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl is required for XLSX metadata extraction")

        metadata: Dict[str, Any] = {
            "file_path": str(file_path),
            "file_type": "office_spreadsheet",
            "file_size": file_path.stat().st_size,
        }

        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True)

            core_props = workbook.properties
            metadata["properties"] = {
                "title": core_props.title,
                "creator": core_props.creator,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "description": core_props.description,
                "category": core_props.category,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "last_modified_by": core_props.lastModifiedBy,
                "revision": core_props.revision,
                "language": core_props.language,
            }

            metadata["workbook_info"] = {
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
            }

            workbook.close()

        except Exception as e:
            logger.warning(f"Error extracting XLSX metadata from {file_path}: {e}")
            metadata["error"] = str(e)

        return metadata

    def extract_metadata(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """Extract metadata from file based on file type.

        Args:
            file_path: Path to file

        Returns:
            Dictionary containing metadata or None if file type not supported
        """
        suffix = file_path.suffix.lower()

        try:
            if suffix in self.IMAGE_EXTENSIONS:
                self.stats["images"] += 1
                return self.extract_image_metadata(file_path)
            elif suffix in self.AUDIO_EXTENSIONS:
                self.stats["audio"] += 1
                return self.extract_audio_metadata(file_path)
            elif suffix == ".docx":
                self.stats["office"] += 1
                return self.extract_docx_metadata(file_path)
            elif suffix == ".xlsx":
                self.stats["office"] += 1
                return self.extract_xlsx_metadata(file_path)
            else:
                logger.debug(f"Unsupported file type: {suffix}")
                return None

        except ImportError as e:
            logger.error(f"Required library not available: {e}")
            self.stats["errors"] += 1
            return None
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            self.stats["errors"] += 1
            return None

    def process_files(
        self, file_paths: List[Path], recursive: bool = False
    ) -> List[Dict[str, Any]]:
        """Process multiple files and extract metadata.

        Args:
            file_paths: List of file paths or directory paths
            recursive: If True, recursively scan directories

        Returns:
            List of metadata dictionaries
        """
        all_metadata: List[Dict[str, Any]] = []
        files_to_process: List[Path] = []

        for path in file_paths:
            path = Path(path).expanduser().resolve()

            if path.is_file():
                files_to_process.append(path)
            elif path.is_dir():
                if recursive:
                    for ext in (
                        self.IMAGE_EXTENSIONS
                        | self.AUDIO_EXTENSIONS
                        | self.OFFICE_EXTENSIONS
                    ):
                        files_to_process.extend(path.rglob(f"*{ext}"))
                else:
                    for ext in (
                        self.IMAGE_EXTENSIONS
                        | self.AUDIO_EXTENSIONS
                        | self.OFFICE_EXTENSIONS
                    ):
                        files_to_process.extend(path.glob(f"*{ext}"))
            else:
                logger.warning(f"Path does not exist: {path}")

        logger.info(f"Found {len(files_to_process)} files to process")

        for file_path in files_to_process:
            self.stats["processed"] += 1
            metadata = self.extract_metadata(file_path)
            if metadata:
                all_metadata.append(metadata)

        return all_metadata

    def format_report(self, metadata_list: List[Dict[str, Any]]) -> str:
        """Format metadata as a readable report.

        Args:
            metadata_list: List of metadata dictionaries

        Returns:
            Formatted string report
        """
        if not metadata_list:
            return "No metadata extracted."

        lines = ["Metadata Extraction Report", "=" * 80, ""]
        lines.append(f"Total files processed: {len(metadata_list)}")
        lines.append(f"Images: {self.stats['images']}")
        lines.append(f"Audio: {self.stats['audio']}")
        lines.append(f"Office files: {self.stats['office']}")
        lines.append(f"Errors: {self.stats['errors']}")
        lines.append("")
        lines.append("-" * 80)

        for metadata in metadata_list:
            lines.append("")
            lines.append(f"File: {metadata.get('file_path', 'Unknown')}")
            lines.append(f"Type: {metadata.get('file_type', 'Unknown')}")
            lines.append(f"Size: {metadata.get('file_size', 0):,} bytes")

            if metadata.get("file_type") == "image":
                if "size" in metadata:
                    lines.append(
                        f"Dimensions: {metadata['size']['width']}x{metadata['size']['height']}"
                    )
                if "exif" in metadata:
                    lines.append("EXIF Data:")
                    for key, value in metadata["exif"].items():
                        if key != "GPSInfo":
                            lines.append(f"  {key}: {value}")

            elif metadata.get("file_type") == "audio":
                if "duration" in metadata:
                    lines.append(f"Duration: {metadata['duration']:.2f} seconds")
                if "tags" in metadata:
                    lines.append("ID3 Tags:")
                    for key, value in metadata["tags"].items():
                        lines.append(f"  {key}: {value}")

            elif metadata.get("file_type") in ["office_document", "office_spreadsheet"]:
                if "properties" in metadata:
                    lines.append("Document Properties:")
                    for key, value in metadata["properties"].items():
                        if value:
                            lines.append(f"  {key}: {value}")

            if "error" in metadata:
                lines.append(f"Error: {metadata['error']}")

            lines.append("-" * 80)

        return "\n".join(lines)


def load_config(config_path: Path) -> dict:
    """Load configuration from YAML file.

    Args:
        config_path: Path to configuration file

    Returns:
        Dictionary containing configuration values

    Raises:
        FileNotFoundError: If config file does not exist
        yaml.YAMLError: If config file is invalid YAML
    """
    if not config_path.exists():
        raise FileNotFoundError(f"Configuration file not found: {config_path}")

    try:
        with open(config_path, "r") as f:
            config = yaml.safe_load(f)
        return config if config else {}
    except yaml.YAMLError as e:
        logger.error(f"Invalid YAML in config file: {e}")
        raise


def main() -> int:
    """Main entry point for the script.

    Returns:
        Exit code (0 for success, non-zero for error)
    """
    parser = argparse.ArgumentParser(
        description="Extract and report file metadata (EXIF, ID3, document properties)"
    )
    parser.add_argument(
        "paths",
        type=str,
        nargs="+",
        help="File paths or directory paths to process",
    )
    parser.add_argument(
        "--recursive",
        action="store_true",
        help="Recursively scan directories",
    )
    parser.add_argument(
        "--output",
        type=str,
        default=None,
        help="Output file path for JSON report",
    )
    parser.add_argument(
        "--report",
        type=str,
        default=None,
        help="Output file path for text report",
    )
    parser.add_argument(
        "--config",
        type=str,
        default=None,
        help="Path to configuration file (YAML)",
    )

    args = parser.parse_args()

    try:
        if args.config:
            config = load_config(Path(args.config))
            if "recursive" in config:
                args.recursive = config["recursive"]

        extractor = MetadataExtractor()
        file_paths = [Path(p) for p in args.paths]
        metadata_list = extractor.process_files(file_paths, recursive=args.recursive)

        if args.output:
            output_path = Path(args.output)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, "w") as f:
                json.dump(metadata_list, f, indent=2, default=str)
            logger.info(f"JSON report saved to {output_path}")

        if args.report:
            report_path = Path(args.report)
            report_path.parent.mkdir(parents=True, exist_ok=True)
            report_text = extractor.format_report(metadata_list)
            with open(report_path, "w") as f:
                f.write(report_text)
            logger.info(f"Text report saved to {report_path}")
        else:
            print(extractor.format_report(metadata_list))

        return 0

    except (ValueError, FileNotFoundError) as e:
        logger.error(f"Error: {e}")
        return 1
    except Exception as e:
        logger.exception(f"Unexpected error: {e}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
